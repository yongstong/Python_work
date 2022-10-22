import openpyxl
from docx import Document

from docx.oxml.ns import qn
from docx.shared import Pt

def choice_fill(choice_num,choice_each_score,doc):
    
    p = doc.add_paragraph()
    text = p.add_run("一、选择题（共{}道，每题{}分）".format(choice_num,choice_each_score))
    text.font.size = Pt(15)                                # 字体大小
    text.font.name = '宋体'                     # 控制是西文时的字体
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    for i in range(int(choice_num)):
        p = doc.add_paragraph()
        text = p.add_run("{}.(题目)".format(i+1))
        text.font.size = Pt(15)
        p = doc.add_paragraph()
        text = p.add_run("A.(选项A)")
        text.font.size = Pt(15)
        p = doc.add_paragraph()
        text = p.add_run("B.(选项B)")
        text.font.size = Pt(15)
        p = doc.add_paragraph()
        text = p.add_run("C.(选项C)")
        text.font.size = Pt(15)
        p = doc.add_paragraph()
        text = p.add_run("D.(选项D)")
        text.font.size = Pt(15)  # 字体大小
        p = doc.add_paragraph("")
        text.font.name = '宋体'                     # 控制是西文时的字体
        text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
        

def choice_replace():
    workbook = openpyxl.load_workbook('question.xlsx')
    sheet = workbook.active
    doc = Document('shijuan.docx')
    
    for j in range(2,sheet.max_row,1):
        cell1 = sheet.cell(row=j, column=4)
        cella = sheet.cell(row=j, column=5)
        cellb = sheet.cell(row=j, column=6)
        cellc = sheet.cell(row=j, column=7)
        celld = sheet.cell(row=j, column=8)
        i = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '(题目)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(题目)', str(cell1.value))
                    i += 1

                elif '(选项A)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(选项A)', str(cella.value))
                    i += 1

                elif '(选项B)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(选项B)', str(cellb.value))
                    i += 1

                elif '(选项C)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(选项C)', str(cellc.value))
                    i += 1

                elif '(选项D)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(选项D)', str(celld.value))
                    i += 1
                if i == 5: break
            if i == 5: break
    doc.save('shijuan.docx')



if __name__ == "__main__":
    doc = Document('2019-2020（B2）试卷.docx')
    choice_num = input("选择题数目：\n")
    choice_each_score = input("每题分数：\n")
    choice_fill(choice_num,choice_each_score,doc)
    doc.save('shijuan.docx')
    choice_replace()
    
    