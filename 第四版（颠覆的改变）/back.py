import openpyxl
from docx import Document
import random
from docx.oxml.ns import qn
from docx.shared import Pt


def choic(choice_num,choice_each_score,begin,num):
    for i in range(int(begin),int(begin)+int(num)):
        choice_fill(choice_num,choice_each_score,i)
        choice_replace(i,choice_num)

def choice_fill(choice_num,choice_each_score,dox_name):
    doc = Document('data/2019-2020（B2）试卷.docx')
    
    p = doc.add_paragraph()
    text = p.add_run("一、选择题（共{}道，每题{}分）".format(choice_num,choice_each_score))
    text.font.size = Pt(15)                                # 字体大小
    text.font.name = '宋体'                     # 控制是西文时的字体
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    
    for i in range(int(choice_num)):
        p = doc.add_paragraph()
        text = p.add_run("{}.(题目)\nA.(选项A)\nB.(选项B)\nC.(选项C)\nD.(选项D)\n".format(i+1))
        text.font.size = Pt(15)
        text.font.name = '宋体'                     # 控制是西文时的字体
        text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    text = p.add_run("\n\n")
    doc.save(f'test/{dox_name}.docx')
        
def choice_replace(dox_name,num):
    workbook = openpyxl.load_workbook('data/choice.xlsx')
    sheet = workbook.active
    doc = Document(f'test/{dox_name}.docx')
    lis = random.sample(range(2,sheet.max_row), int(num))
    for j in lis:
        cell1 = sheet.cell(row=j, column=4)
        cella = sheet.cell(row=j, column=5)
        cellb = sheet.cell(row=j, column=6)
        cellc = sheet.cell(row=j, column=7)
        celld = sheet.cell(row=j, column=8)
        sign = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '(题目)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(题目)', str(cell1.value))
                    sign += 1

                if '(选项A)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(选项A)', str(cella.value))
                    sign += 1

                if '(选项B)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(选项B)', str(cellb.value))
                    sign += 1

                if '(选项C)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(选项C)', str(cellc.value))
                    sign += 1

                if '(选项D)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(选项D)', str(celld.value))
                    sign += 1
                if sign == 5:
                    break
            if sign == 5: break
    doc.save(f'test/{dox_name}.docx')

#--------------------------------------------

def fill(fill_num,fill_each_score,begin,num):
    for i in range(int(begin),int(begin)+int(num)):
        fill_fill(fill_num,fill_each_score,i)
        fill_replace(i,fill_num)

def fill_fill(fill_num,fill_each_score,dox_name):
    doc = Document(f'test/{dox_name}.docx')
    
    p = doc.add_paragraph()
    text = p.add_run("二、选择题（共{}道，每题{}分）".format(fill_num,fill_each_score))
    text.font.size = Pt(15)                                # 字体大小
    text.font.name = '宋体'                     # 控制是西文时的字体
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    
    for i in range(int(fill_num)):
        p = doc.add_paragraph()
        text = p.add_run("{}.(题目)\n".format(i+1))
        text.font.size = Pt(15)
        text.font.name = '宋体'                     # 控制是西文时的字体
        text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    text = p.add_run("\n\n")
    doc.save(f'test/{dox_name}.docx')
        
def fill_replace(dox_name,num):
    workbook = openpyxl.load_workbook('data/fill.xlsx')
    sheet = workbook.active
    doc = Document(f'test/{dox_name}.docx')
    lis = random.sample(range(2,sheet.max_row), int(num))
    for j in lis:
        cell1 = sheet.cell(row=j, column=3)
        sign = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '(题目)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(题目)', str(cell1.value))
                    sign += 1

                if sign == 1:
                    break
            if sign == 1: break
    doc.save(f'test/{dox_name}.docx')
    
#-----------------------------------------  
    
def answer(answer_num,answer_each_score,begin,num):
    for i in range(int(begin),int(begin)+int(num)):
        answer_fill(answer_num,answer_each_score,i)
        answer_replace(i,answer_num)

def answer_fill(answer_num,answer_each_score,dox_name):
    doc = Document(f'test/{dox_name}.docx')
    
    p = doc.add_paragraph()
    text = p.add_run("三、解答题（共{}道，每题{}分）".format(answer_num,answer_each_score))
    text.font.size = Pt(15)                                # 字体大小
    text.font.name = '宋体'                     # 控制是西文时的字体
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    
    for i in range(int(answer_num)):
        p = doc.add_paragraph()
        text = p.add_run("{}.(题目)\n".format(i+1))
        text.font.size = Pt(15)
        text.font.name = '宋体'                     # 控制是西文时的字体
        text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    text = p.add_run("\n\n")
    doc.save(f'test/{dox_name}.docx')
        
def answer_replace(dox_name,num):
    workbook = openpyxl.load_workbook('data/answer.xlsx')
    sheet = workbook.active
    doc = Document(f'test/{dox_name}.docx')
    lis = random.sample(range(2,sheet.max_row), int(num))
    for j in lis:
        cell1 = sheet.cell(row=j, column=3)
        sign = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '(题目)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(题目)', str(cell1.value))
                    sign += 1

                if sign == 1:
                    break
            if sign == 1: break
    doc.save(f'test/{dox_name}.docx')   
    
#----------------------------  
    
def calculte(calculte_num,calculte_each_score,begin,num):
    for i in range(int(begin),int(begin)+int(num)):
        calculte_fill(calculte_num,calculte_each_score,i)
        calculte_replace(i,calculte_num)

def calculte_fill(calculte_num,calculte_each_score,dox_name):
    doc = Document(f'test/{dox_name}.docx')
    
    p = doc.add_paragraph()
    text = p.add_run("三、解答题（共{}道，每题{}分）".format(calculte_num,calculte_each_score))
    text.font.size = Pt(15)                                # 字体大小
    text.font.name = '宋体'                     # 控制是西文时的字体
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    
    for i in range(int(calculte_num)):
        p = doc.add_paragraph()
        text = p.add_run("{}.(题目)\n".format(i+1))
        text.font.size = Pt(15)
        text.font.name = '宋体'                     # 控制是西文时的字体
        text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 控制是中文时的字体
    text = p.add_run("\n\n")
    doc.save(f'test/{dox_name}.docx')
        
def calculte_replace(dox_name,num):
    workbook = openpyxl.load_workbook('data/calculte.xlsx')
    sheet = workbook.active
    doc = Document(f'test/{dox_name}.docx')
    lis = random.sample(range(2,sheet.max_row), int(num))
    for j in lis:
        cell1 = sheet.cell(row=j, column=3)
        sign = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '(题目)' in run.text:
					# 如果存在匹配得字符串，那么将当前得run替换成合并后得字符串
                    run.text = run.text.replace('(题目)', str(cell1.value))
                    sign += 1

                if sign == 1:
                    break
            if sign == 1: break
    doc.save(f'test/{dox_name}.docx')   
 
if __name__ == '__main__':
    choice_num=input("shu")
    choice_each_score=input("fen")
    choic(choice_num,choice_each_score)